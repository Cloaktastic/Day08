import os
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

LEGAL_DIR = Path("data/landing/legal")
LEGAL_DIR.mkdir(parents=True, exist_ok=True)

def build_pdf(filename, title, content):
    filepath = LEGAL_DIR / filename
    doc = SimpleDocTemplate(str(filepath), pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    
    # Custom styles to support unicode (default Helvetica doesn't support Vietnamese well, but for PDF text extractor/markitdown it's fine as long as the text is encoded)
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        alignment=TA_CENTER,
        spaceAfter=20
    )
    
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['BodyText'],
        alignment=TA_JUSTIFY,
        spaceAfter=10
    )
    
    story = [Paragraph(title, title_style), Spacer(1, 12)]
    
    for section_title, paragraphs in content.items():
        heading_style = ParagraphStyle(
            'HeadingStyle',
            parent=styles['Heading2'],
            spaceBefore=12,
            spaceAfter=6
        )
        story.append(Paragraph(section_title, heading_style))
        for p in paragraphs:
            story.append(Paragraph(p, body_style))
            
    doc.build(story)
    print(f"Created PDF: {filepath}")

def build_docx(filename, title, content):
    filepath = LEGAL_DIR / filename
    doc = Document()
    doc.add_heading(title, 0)
    
    for section_title, paragraphs in content.items():
        doc.add_heading(section_title, level=1)
        for p in paragraphs:
            doc.add_paragraph(p)
            
    doc.save(filepath)
    print(f"Created DOCX: {filepath}")

# Content definitions
luat_content = {
    "Chương I: Quy định chung": [
        "Điều 1. Phạm vi điều chỉnh: Luật này quy định về phòng ngừa, ngăn chặn và đấu tranh chống tội phạm và tệ nạn ma túy; kiểm soát các hoạt động hợp pháp liên quan đến ma túy; trách nhiệm của cá nhân, gia đình, cơ quan, tổ chức trong phòng, chống ma túy; cai nghiện ma túy và quản lý nhà nước về phòng, chống ma túy.",
        "Điều 2. Giải thích từ ngữ: Trong Luật này, các từ ngữ dưới đây được hiểu như sau:\n"
        "1. Chất ma túy là chất gây nghiện, chất hướng thần được quy định trong danh mục chất ma túy do Chính phủ ban hành.\n"
        "2. Tiền chất là hóa chất không thể thiếu được trong quá trình điều chế, sản xuất chất ma túy được quy định trong danh mục do Chính phủ ban hành.\n"
        "3. Người sử dụng trái phép chất ma túy là người có hành vi tự ý sử dụng chất ma túy trái quy định của pháp luật và bị cơ quan có thẩm quyền xét nghiệm phát hiện chất ma túy trong cơ thể.\n"
        "4. Người nghiện ma túy là người sử dụng chất ma túy, thuốc gây nghiện, thuốc hướng thần và bị lệ thuộc vào các chất này."
    ],
    "Chương II: Trách nhiệm phòng, chống ma túy": [
        "Điều 3. Chính sách của Nhà nước về phòng, chống ma túy:\n"
        "1. Thực hiện đồng bộ các biện pháp phòng, chống ma túy; kết hợp giữa phòng ngừa với đấu tranh chống tội phạm về ma túy; coi trọng công tác tuyên truyền, giáo dục, kết hợp giữa gia đình, nhà trường và xã hội.\n"
        "2. Khuyến khích cá nhân, gia đình, cơ quan, tổ chức tham gia phòng, chống ma túy; hỗ trợ hoạt động cai nghiện ma túy tự nguyện.",
        "Điều 4. Trách nhiệm của cá nhân, gia đình:\n"
        "1. Tuyên truyền, giáo dục thành viên trong gia đình về tác hại của ma túy và chấp hành quy định của pháp luật về phòng, chống ma túy.\n"
        "2. Đấu tranh phòng, chống ma túy tại cộng đồng dân cư, thông báo kịp thời cho cơ quan chức năng khi phát hiện hành vi vi phạm pháp luật về ma túy."
    ]
}

nghi_dinh_content = {
    "Chương I: Công tác phối hợp đấu tranh chống tội phạm về ma túy": [
        "Điều 1. Nguyên tắc phối hợp: Công tác phối hợp giữa các cơ quan chuyên trách phòng, chống tội phạm về ma túy phải được thực hiện chủ động, kịp thời, đúng chức năng, nhiệm vụ, quyền hạn do pháp luật quy định và dưới sự chỉ đạo tập trung, thống nhất của Chính phủ.",
        "Điều 2. Các cơ quan chuyên trách phối hợp:\n"
        "1. Lực lượng Cảnh sát điều tra tội phạm về ma túy thuộc Bộ Công an.\n"
        "2. Lực lượng chuyên trách phòng, chống tội phạm ma túy thuộc Bộ đội Biên phòng, Cảnh sát biển và lực lượng Hải quan."
    ],
    "Chương II: Kiểm soát các hoạt động hợp pháp liên quan đến ma túy": [
        "Điều 5. Quản lý, kiểm soát tiền chất: Các cơ quan, tổ chức có hoạt động xuất khẩu, nhập khẩu, tạm nhập, tái xuất, vận chuyển tiền chất phải thực hiện khai báo, cấp phép theo đúng quy định. Nghiêm cấm mọi hành vi sử dụng tiền chất sai mục đích để điều chế chất ma túy trái phép.",
        "Điều 6. Xử lý vi phạm trong việc mua bán, vận chuyển tiền chất trái phép: Các trường hợp vi phạm sẽ bị xử lý hành chính hoặc hình sự tùy thuộc vào tính chất, mức độ của hành vi vi phạm."
    ]
}

bo_luat_hinh_su_content = {
    "Chương XX: Các tội phạm về ma túy": [
        "Điều 248. Tội sản xuất trái phép chất ma túy:\n"
        "1. Người nào sản xuất trái phép chất ma túy dưới bất kỳ hình thức nào, thì bị phạt tù từ 02 năm đến 07 năm.\n"
        "2. Phạm tội thuộc một trong các trường hợp nghiêm trọng (có tổ chức, tái phạm nguy hiểm, số lượng lớn) thì bị phạt tù từ 07 năm đến 15 năm, hoặc phạt tù từ 15 năm đến 20 năm, tù chung thân hoặc tử hình.",
        "Điều 249. Tội tàng trữ trái phép chất ma túy:\n"
        "1. Người nào tàng trữ trái phép chất ma túy mà không nhằm mục đích mua bán, vận chuyển, sản xuất trái phép chất ma túy thuộc một trong các trường hợp quy định thì bị phạt tù từ 01 năm đến 05 năm.\n"
        "2. Phạm tội tàng trữ số lượng lớn, tái phạm nguy hiểm thì bị phạt tù từ 05 năm đến 10 năm, 10 năm đến 15 năm, hoặc từ 15 năm đến 20 năm hoặc tù chung thân.",
        "Điều 250. Tội vận chuyển trái phép chất ma túy:\n"
        "1. Người nào vận chuyển trái phép chất ma túy mà không nhằm mục đích sản xuất, mua bán, tàng trữ trái phép chất ma túy, thì bị phạt tù từ 02 năm đến 07 năm.\n"
        "2. Phạm tội vận chuyển số lượng đặc biệt lớn thì bị phạt tù 20 năm, tù chung thân hoặc tử hình.",
        "Điều 251. Tội mua bán trái phép chất ma túy:\n"
        "1. Người nào mua bán trái phép chất ma túy, thì bị phạt tù từ 02 năm đến 07 năm.\n"
        "2. Phạm tội mua bán chất ma túy có tổ chức, chuyên nghiệp hoặc số lượng lớn thì phạt tù từ 07 năm đến 15 năm, 15 năm đến 20 năm, tù chung thân hoặc tử hình."
    ]
}

# Generate 3 files
build_pdf("luat-phong-chong-ma-tuy-2021.pdf", "LUẬT PHÒNG, CHỐNG MA TÚY 2021", luat_content)
build_pdf("nghi-dinh-105-2021.pdf", "NGHỊ ĐỊNH 105/2021/NĐ-CP HƯỚNG DẪN THI HÀNH LUẬT PHÒNG CHỐNG MA TÚY", nghi_dinh_content)
build_docx("bo-luat-hinh-su-2015.docx", "BỘ LUẬT HÌNH SỰ 2015 (CHƯƠNG XX: CÁC TỘI PHẠM VỀ MA TÚY)", bo_luat_hinh_su_content)
